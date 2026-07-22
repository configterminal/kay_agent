#!/usr/bin/env python3
"""
【遗留】课程 .docx → Markdown（无时间戳）。

新录入源应为「视频转写的带时间戳 md」。本脚本仅用于过渡期把旧字幕/讲义
转成可读 md，**不要**把产出直接当作新 RAG 索引唯一源。

有对应 .mp4 时请改用 transcribe_video.py。

用法:
  python skills/course-transcribe/scripts/convert_docx.py --dry-run
  python skills/course-transcribe/scripts/convert_docx.py --course RAG101
"""

from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path

try:
    import docx
except ImportError:
    sys.exit("请先安装 python-docx: pip install python-docx --target f:\\jupyter")

from auto_fix_rules import apply_fix_and_flag


def extract_text(doc: "docx.Document") -> str:
    """拼接全部段落与表格文字。"""
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def is_heading_like(line: str) -> tuple[bool, int]:
    """粗判短行是否像幻灯标题 → Markdown 标题级别。"""
    stripped = line.strip()
    if not stripped:
        return False, 0
    if re.match(r"^#{1,6}\s", stripped):
        return True, 0
    if len(stripped) <= 60 and not re.search(
        r"[。，、；：？！…\.\,\;\!\?]$", stripped
    ):
        if re.match(r"^\d{1,2}[-–—]\d{1,2}\s+\S", stripped):
            return True, 2
        if re.match(r"^第[一二三四五六七八九十\d]+章", stripped):
            return True, 1
        if len(stripped) <= 25:
            return True, 3
    return False, 0


def clean_title(line: str) -> str:
    """去掉标题里的 _ev 后缀。"""
    return re.sub(r"_ev\b", "", line).strip()


def docx_to_markdown(docx_path: Path, flag: bool = True) -> str:
    """单文件 docx → md 字符串（无时间戳）。"""
    doc = docx.Document(str(docx_path))
    raw = extract_text(doc)
    out: list[str] = []

    for line in raw.split("\n"):
        stripped = line.strip()
        if not stripped:
            if out and out[-1] != "":
                out.append("")
            continue

        is_heading, level = is_heading_like(stripped)
        if is_heading and level > 0:
            out.append("")
            out.append(f"{'#' * level} {clean_title(stripped)}")
            out.append("")
            continue

        body = apply_fix_and_flag(stripped) if flag else stripped
        out.append(body)

    result = "\n".join(out)
    result = re.sub(r"\n_ev\b", "", result)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip() + "\n"


def find_docx_files(base: Path, course_filter: str | None = None) -> list[Path]:
    files: list[Path] = []
    for root, _dirs, filenames in os.walk(base.resolve()):
        for fn in filenames:
            if fn.endswith(".docx") and not fn.startswith("~$"):
                fp = Path(root) / fn
                if course_filter and course_filter not in str(fp):
                    continue
                files.append(fp)
    return sorted(files)


def main() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

    ap = argparse.ArgumentParser(
        description="【遗留】课程 docx → md（无时间戳；有视频请用 transcribe_video.py）"
    )
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--course", type=str, default=None)
    ap.add_argument("--no-flag", action="store_true", help="不插入 ⚠️ 审核注释")
    args = ap.parse_args()

    repo = Path(__file__).resolve().parents[3]
    courses_root = repo / "resources" / "courses"
    if not courses_root.is_dir():
        sys.exit(f"courses 目录不存在: {courses_root}")

    print(
        "注意: convert_docx 为遗留路径，产出无 [时间戳]。"
        "新索引请以视频转写 md 为准。\n"
    )

    files = find_docx_files(courses_root, args.course)
    if not files:
        print("没有找到 .docx 文件。")
        return

    print(f"找到 {len(files)} 个 docx 文件:\n")
    for fp in files:
        print(f"  {fp.relative_to(repo)}")

    if args.dry_run:
        print("\n[dry-run] 预览前 300 字符:\n")
        for fp in files:
            md = docx_to_markdown(fp, flag=not args.no_flag)
            print(f"--- {fp.name} ---")
            print(md[:300])
            print("...\n")
        return

    print("\n开始转换...\n")
    ok = fail = 0
    for fp in files:
        try:
            md = docx_to_markdown(fp, flag=not args.no_flag)
            out_path = fp.with_suffix(".md")
            out_path.write_text(md, encoding="utf-8")
            print(f"  [OK] {out_path.relative_to(repo)}")
            ok += 1
        except Exception as e:
            print(f"  [FAIL] {fp.name}: {e}")
            fail += 1

    print(f"\n完成: {ok} 个成功, {fail} 个失败")


if __name__ == "__main__":
    main()

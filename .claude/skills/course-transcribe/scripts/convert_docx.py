#!/usr/bin/env python3
"""
Course docx → Markdown converter.

Usage:
  python convert_docx.py                          # convert ALL docx under resources/courses/
  python convert_docx.py --dry-run                # preview only, no files written
  python convert_docx.py --course "RAG101"       # filter by course dir

What it does per .docx:
  1. Extracts all paragraphs + table text.
  2. Detects slide-like headings and promotes them to Markdown `##` / `###`.
  3. Outputs a clean .md alongside the .docx (same chapter dir).

Call from repo root: `f:\agent`.
"""
import argparse, json, os, re, sys
from pathlib import Path

try:
    import docx
except ImportError:
    sys.exit("请先安装 python-docx: pip install python-docx --target f:\\jupyter")

# --------------- ASR error correction ---------------
# HIGH confidence — auto-fix directly (no markers)
AUTO_FIX: list[tuple[str, str]] = [
    # English terms — use (?<![a-zA-Z])…(?![a-zA-Z]) for CJK-friendly boundary
    (r'(?<![a-zA-Z])rig(?![a-zA-Z])', "RAG"),
    (r'(?<![a-zA-Z])reg(?![a-zA-Z])', "RAG"),
    (r'(?<![a-zA-Z])ig(?![a-zA-Z])', "RAG"),
    (r'OPenAi', "OpenAI"),
    (r'deep\s*seek', "DeepSeek"),
    (r'\blamer\b', "LLaMA"),
    (r'LA\s*mer', "LLaMA"),
    (r'line\s*chain', "LangChain"),
    (r'p\s*touch', "PyTorch"),
    (r'g\s*radio', "Gradio"),
    (r'near\s*for\s*z', "Neo4j"),
    (r'jupiter\s*note', "Jupyter Notebook"),
    (r'jupiter', "Jupyter"),
    (r"ra'g", "RAG"),
    # Chinese terms
    (r'大圆模型', "大语言模型"),
    (r'大元模型', "大语言模型"),
    (r'大于模型', "大语言模型"),
    (r'单元模型', "大语言模型"),
    (r'代元模型', "大语言模型"),
    (r'大圆模', "大语言模"),
    (r'掐gpt', "ChatGPT"),
]

# LOW confidence — flag with HTML comment for human review
FLAG_ONLY: list[tuple[str, str]] = [
    (r'dream\s*nine', "dream nine → Gemini?"),
    (r'dream\s*1\s*[.]?\s*5', "dream 1.5 → Gemini 1.5?"),
    (r'dream\s*2\s*[.]?\s*0', "dream 2.0 → Gemini 2.0?"),
    (r'km\s*it', "km it → Kimi?"),
    (r'\bg\s*bd\b', "g bd → GPT?"),
    (r'\bg\s*b[dt]\b', "gbt/gbd → GPT?"),
    (r'gha\s*PC', "gha PC → ChatGPT?"),
    (r'opal\s*on', "opal on → OpenAI?"),
    (r'(?<![a-zA-Z])cord\s*3(?![a-zA-Z])', "cord 3 → Cohere / Claude 3?"),
    (r'(?<![a-zA-Z.])com(?![a-zA-Z])', "com → Chroma?"),
    (r'(?<![a-zA-Z])rap(?![a-zA-Z])', "rap → RAP?"),
    (r'd\s*p\s*c\s*r\s*e', "d p c r e → ?"),
]
# ---------------------------------------------------


def extract_text(doc: "docx.Document") -> str:
    """Return concatenated text of all paragraphs + tables."""
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def is_heading_like(line: str) -> tuple[bool, int]:
    """Guess whether a short standalone line is a slide title → heading level."""
    stripped = line.strip()
    if not stripped:
        return False, 0
    # Already a markdown heading
    if re.match(r'^#{1,6}\s', stripped):
        return True, 0
    # Short enough to be a title, not ending with sentence punctuation
    if len(stripped) <= 60 and not re.search(r'[。，、；：？！…\.\,\;\!\?]$', stripped):
        # Lines like "2-3 解锁RAG三大核心_ev"
        if re.match(r'^\d+[-–—]\d+\s+\S', stripped):
            return True, 2  # ##
        # Lines like "第X章 …"
        if re.match(r'^第[一二三四五六七八九十\d]+章', stripped):
            return True, 1  # #
        # Other short potential headings
        if len(stripped) <= 25:
            return True, 3  # ###
    return False, 0


def clean_title(line: str) -> str:
    """Strip the _ev suffix from section titles."""
    return re.sub(r'_ev\b', '', line).strip()


def fix_and_flag(text: str) -> str:
    """Auto-fix high-confidence errors, flag low-confidence ones with <!--⚠️ -->."""
    # 1. Auto-fix first
    for pattern, replacement in AUTO_FIX:
        text = re.sub(pattern, replacement, text)
    # 2. Flag remaining
    for pattern, note in FLAG_ONLY:
        def _repl(m: re.Match, _note: str = note) -> str:
            return f"{m.group()}<!--⚠️ {_note}-->"
        text = re.sub(pattern, _repl, text)
    return text


def docx_to_markdown(docx_path: Path) -> str:
    """Convert a single .docx to markdown string."""
    doc = docx.Document(str(docx_path))
    raw = extract_text(doc)

    lines = raw.split("\n")
    out: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if out and out[-1] != "":
                out.append("")
            continue

        is_heading, level = is_heading_like(stripped)
        if is_heading and level > 0:
            out.append("")  # blank line before heading
            prefix = "#" * level
            clean = clean_title(stripped)
            out.append(f"{prefix} {clean}")
            out.append("")
            continue

        # Check if it looks like a continuation (no period, short-ish)
        # We keep it as body text
        out.append(fix_and_flag(stripped))

    result = "\n".join(out)

    # Remove standalone footnote artifacts
    result = re.sub(r'\n_ev\b', '', result)

    # Collapse 3+ blank lines into 2
    result = re.sub(r'\n{3,}', '\n\n', result)

    return result.strip() + "\n"


def find_docx_files(base: Path, course_filter: str | None = None) -> list[Path]:
    """Yield all .docx paths under base, optionally filtered by course dir name."""
    base = base.resolve()
    files: list[Path] = []
    for root, _dirs, filenames in os.walk(base):
        for fn in filenames:
            if fn.endswith(".docx") and not fn.startswith("~$"):
                fp = Path(root) / fn
                if course_filter and course_filter not in str(fp):
                    continue
                files.append(fp)
    return sorted(files)


def main() -> None:
    ap = argparse.ArgumentParser(description="Convert course .docx to .md")
    ap.add_argument("--dry-run", action="store_true", help="Preview only")
    ap.add_argument("--course", type=str, default=None, help="Filter by course dir (e.g. RAG101)")
    ap.add_argument("--no-flag", action="store_true", help="Skip error-flag comments")
    args = ap.parse_args()

    global FLAG_ONLY
    if args.no_flag:
        FLAG_ONLY = []

    repo = Path(__file__).resolve().parent.parent.parent.parent  # f:\agent
    courses_root = repo / "resources" / "courses"

    if not courses_root.is_dir():
        sys.exit(f"courses 目录不存在: {courses_root}")

    files = find_docx_files(courses_root, args.course)

    if not files:
        print("没有找到 .docx 文件。")
        return

    print(f"找到 {len(files)} 个 docx 文件:\n")

    for fp in files:
        rel = fp.relative_to(repo)
        rel_str = str(rel)
        print(f"  {rel_str}")

    if args.dry_run:
        print("\n[dry-run] 以下为预览（只显示前 300 字符）:\n")
        for fp in files:
            md = docx_to_markdown(fp)
            print(f"--- {fp.name} ---")
            print(md[:300])
            print("...\n")
        return

    print("\n开始转换...\n")
    ok = fail = 0
    for fp in files:
        try:
            md = docx_to_markdown(fp)
            out_path = fp.with_suffix(".md")
            out_path.write_text(md, encoding="utf-8")
            rel = out_path.relative_to(repo)
            print(f"  [OK] {rel}")
            ok += 1
        except Exception as e:
            print(f"  [FAIL] {fp.name}: {e}")
            fail += 1

    print(f"\n完成: {ok} 个成功, {fail} 个失败")


if __name__ == "__main__":
    main()
